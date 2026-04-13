# api/app/routers/reportes.py
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import func
from io import BytesIO
from datetime import datetime

from app.data.db import get_db
from app.data.models import Autoparte, Pedido, DetallePedido, Usuario

router = APIRouter(
    prefix="/v1/reportes",
    tags=["Reportes"]
)

# ── Helpers ────────────────────────────────────────────────────
def _datos_inventario(db):
    return db.query(Autoparte).all()

def _datos_bajo_stock(db):
    return db.query(Autoparte).filter(Autoparte.stock < 10).all()

def _datos_ventas(db):
    return db.query(Pedido).all()

def _datos_top_productos(db):
    return (
        db.query(Autoparte.nombre, func.sum(DetallePedido.cantidad).label("total_vendido"))
        .join(DetallePedido, Autoparte.id == DetallePedido.autoparte_id)
        .group_by(Autoparte.nombre)
        .order_by(func.sum(DetallePedido.cantidad).desc())
        .limit(10)
        .all()
    )


# ══════════════════════════════════════════════════════════════
# PDF
# ══════════════════════════════════════════════════════════════
def _generar_pdf(titulo: str, encabezados: list, filas: list) -> BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    elementos = []

    elementos.append(Paragraph(titulo, styles["Title"]))
    elementos.append(Paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]))
    elementos.append(Spacer(1, 12))

    data = [encabezados] + filas
    tabla = Table(data)
    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A1A2A")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F7")]),
        ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("PADDING",    (0, 0), (-1, -1), 6),
    ]))
    elementos.append(tabla)
    doc.build(elementos)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════
# XLSX
# ══════════════════════════════════════════════════════════════
def _generar_xlsx(titulo: str, encabezados: list, filas: list) -> BytesIO:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = titulo[:30]

    ws.append([titulo])
    ws.append([f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"])
    ws.append([])
    ws.append(encabezados)

    header_row = 4
    for col_num, _ in enumerate(encabezados, 1):
        cell = ws.cell(row=header_row, column=col_num)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1A1A2A")
        cell.alignment = Alignment(horizontal="center")

    for fila in filas:
        ws.append(fila)

    for col in ws.columns:
        max_len = max(len(str(c.value or "")) for c in col)
        ws.column_dimensions[col[0].column_letter].width = max(max_len + 4, 12)

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════
# DOCX
# ══════════════════════════════════════════════════════════════
def _generar_docx(titulo: str, encabezados: list, filas: list) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(titulo, 0)
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = "Table Grid"

    # Encabezados
    hdr = tabla.rows[0].cells
    for i, enc in enumerate(encabezados):
        hdr[i].text = enc
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas
    for fila in filas:
        row = tabla.add_row().cells
        for i, val in enumerate(fila):
            row[i].text = str(val)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════
# ENDPOINTS — 4 tipos de reporte × 3 formatos = 12 endpoints
# ══════════════════════════════════════════════════════════════

# ── 1. Inventario completo ─────────────────────────────────────
def _enc_inv(): return ["ID", "Nombre", "Categoría", "Precio", "Stock", "Descripción"]
def _fil_inv(data): return [[a.id, a.nombre, a.categoria, f"${a.precio:.2f}", a.stock, a.descripcion or ""] for a in data]

@router.get("/inventario/pdf")
async def reporte_inventario_pdf(db: Session = Depends(get_db)):
    buf = _generar_pdf("Reporte de Inventario - MACUIN", _enc_inv(), _fil_inv(_datos_inventario(db)))
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=inventario.pdf"})

@router.get("/inventario/xlsx")
async def reporte_inventario_xlsx(db: Session = Depends(get_db)):
    buf = _generar_xlsx("Inventario", _enc_inv(), _fil_inv(_datos_inventario(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=inventario.xlsx"})

@router.get("/inventario/docx")
async def reporte_inventario_docx(db: Session = Depends(get_db)):
    buf = _generar_docx("Reporte de Inventario - MACUIN", _enc_inv(), _fil_inv(_datos_inventario(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=inventario.docx"})


# ── 2. Stock bajo (< 10 unidades) ─────────────────────────────
def _enc_stock(): return ["ID", "Nombre", "Categoría", "Precio", "Stock Actual"]
def _fil_stock(data): return [[a.id, a.nombre, a.categoria, f"${a.precio:.2f}", a.stock] for a in data]

@router.get("/bajo-stock/pdf")
async def reporte_bajo_stock_pdf(db: Session = Depends(get_db)):
    buf = _generar_pdf("Reporte de Stock Bajo - MACUIN", _enc_stock(), _fil_stock(_datos_bajo_stock(db)))
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=bajo_stock.pdf"})

@router.get("/bajo-stock/xlsx")
async def reporte_bajo_stock_xlsx(db: Session = Depends(get_db)):
    buf = _generar_xlsx("Stock Bajo", _enc_stock(), _fil_stock(_datos_bajo_stock(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=bajo_stock.xlsx"})

@router.get("/bajo-stock/docx")
async def reporte_bajo_stock_docx(db: Session = Depends(get_db)):
    buf = _generar_docx("Reporte de Stock Bajo - MACUIN", _enc_stock(), _fil_stock(_datos_bajo_stock(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=bajo_stock.docx"})


# ── 3. Ventas / Pedidos ────────────────────────────────────────
def _enc_ventas(): return ["ID Pedido", "Usuario ID", "Fecha", "Estado", "Total"]
def _fil_ventas(data): return [[p.id, p.usuario_id, p.fecha.strftime("%d/%m/%Y"), p.estado, f"${p.total:.2f}"] for p in data]

@router.get("/ventas/pdf")
async def reporte_ventas_pdf(db: Session = Depends(get_db)):
    buf = _generar_pdf("Reporte de Ventas - MACUIN", _enc_ventas(), _fil_ventas(_datos_ventas(db)))
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=ventas.pdf"})

@router.get("/ventas/xlsx")
async def reporte_ventas_xlsx(db: Session = Depends(get_db)):
    buf = _generar_xlsx("Ventas", _enc_ventas(), _fil_ventas(_datos_ventas(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=ventas.xlsx"})

@router.get("/ventas/docx")
async def reporte_ventas_docx(db: Session = Depends(get_db)):
    buf = _generar_docx("Reporte de Ventas - MACUIN", _enc_ventas(), _fil_ventas(_datos_ventas(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=ventas.docx"})


# ── 4. Top productos más vendidos ──────────────────────────────
def _enc_top(): return ["Autoparte", "Total Vendido"]
def _fil_top(data): return [[nombre, total] for nombre, total in data]

@router.get("/top-productos/pdf")
async def reporte_top_pdf(db: Session = Depends(get_db)):
    buf = _generar_pdf("Top Productos Más Vendidos - MACUIN", _enc_top(), _fil_top(_datos_top_productos(db)))
    return StreamingResponse(buf, media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=top_productos.pdf"})

@router.get("/top-productos/xlsx")
async def reporte_top_xlsx(db: Session = Depends(get_db)):
    buf = _generar_xlsx("Top Productos", _enc_top(), _fil_top(_datos_top_productos(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=top_productos.xlsx"})

@router.get("/top-productos/docx")
async def reporte_top_docx(db: Session = Depends(get_db)):
    buf = _generar_docx("Top Productos Más Vendidos - MACUIN", _enc_top(), _fil_top(_datos_top_productos(db)))
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=top_productos.docx"})
